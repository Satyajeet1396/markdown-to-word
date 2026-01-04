import streamlit as st
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import subprocess
import tempfile
import os

# Page configuration
st.set_page_config(
    page_title="Markdown to Word Converter",
    page_icon="📄",
    layout="wide"
)

# Title and description
st.title("📄 Markdown to Word Converter")
st.markdown("Convert markdown with **editable equations** (Pandoc) or **Unicode symbols** (python-docx)")

# Check Pandoc availability
def check_pandoc():
    """Check if Pandoc is installed"""
    try:
        result = subprocess.run(['pandoc', '--version'], 
                              capture_output=True, 
                              text=True, 
                              timeout=5)
        return result.returncode == 0
    except:
        return False

PANDOC_AVAILABLE = check_pandoc()

# Sidebar for settings
with st.sidebar:
    st.header("🔧 Settings")
    
    # Conversion method
    st.subheader("Conversion Method")
    if PANDOC_AVAILABLE:
        conversion_method = st.radio(
            "Choose method:",
            ["Pandoc (Editable Equations)", "Python-docx (Unicode)"],
            help="Pandoc creates editable Word equations. Python-docx uses Unicode symbols."
        )
        use_pandoc = conversion_method.startswith("Pandoc")
    else:
        st.warning("⚠️ Pandoc not installed - using Python-docx method")
        use_pandoc = False
    
    st.divider()
    
    # GitHub section
    st.subheader("GitHub Repository")
    github_url = st.text_input(
        "GitHub File URL",
        placeholder="https://github.com/user/repo/blob/main/file.md",
        help="Enter the GitHub URL of a markdown file"
    )
    
    if st.button("📥 Fetch from GitHub"):
        if github_url:
            try:
                raw_url = github_url.replace("github.com", "raw.githubusercontent.com").replace("/blob/", "/")
                response = requests.get(raw_url)
                response.raise_for_status()
                st.session_state['markdown_content'] = response.text
                st.success("✅ Successfully fetched from GitHub!")
            except Exception as e:
                st.error(f"❌ Error fetching from GitHub: {str(e)}")
        else:
            st.warning("Please enter a GitHub URL")
    
    st.divider()
    
    # Document settings
    st.subheader("Document Settings")
    doc_title = st.text_input("Document Title", value="Converted Document")
    font_size = st.slider("Base Font Size", 8, 16, 11)
    use_colors = st.checkbox("Use colored headings", value=True)
    
    if use_pandoc:
        st.subheader("Pandoc Options")
        use_toc = st.checkbox("Include Table of Contents", value=False)
        number_sections = st.checkbox("Number sections", value=False)
    
    st.divider()
    
    # Platform info
    st.info(f"""
    **Platform Status:**
    - Pandoc: {'✅ Installed' if PANDOC_AVAILABLE else '❌ Not Available'}
    - Python-docx: ✅ Always Available
    
    **Deployment:**
    - Streamlit Cloud: Use Python-docx
    - GitHub Codespaces: Install Pandoc
    - Local: Both methods work
    """)

# Initialize session state
if 'markdown_content' not in st.session_state:
    st.session_state['markdown_content'] = """# Sample Markdown

## Introduction
This is a **sample** markdown document with *formatting*.

### Math Examples

**For Pandoc:** Use `$...$` syntax
Inline math: $\\alpha$ and $\\beta$

Display math:
$$
S = \\alpha + \\beta
$$

**For Python-docx:** Use `\\(...\\)` syntax
Inline: \\( \\alpha \\) and \\( \\beta \\)

Display:
\\[
E = mc^2
\\]

### Lists and Tables

- Bullet point 1
- Bullet point 2

| Column 1 | Column 2 |
|----------|----------|
| Data 1   | Data 2   |
"""

# PANDOC CONVERSION FUNCTION
def convert_with_pandoc(markdown_text, title, use_toc, number_sections):
    """Convert markdown to Word using Pandoc"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as md_file:
        yaml_header = f"""---
title: "{title}"
---

"""
        md_file.write(yaml_header + markdown_text)
        md_path = md_file.name
    
    output_path = tempfile.mktemp(suffix='.docx')
    
    try:
        cmd = [
            'pandoc',
            md_path,
            '-o', output_path,
            '--from', 'markdown',
            '--to', 'docx',
            '--standalone'
        ]
        
        if use_toc:
            cmd.append('--toc')
        if number_sections:
            cmd.append('--number-sections')
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode != 0:
            raise Exception(f"Pandoc error: {result.stderr}")
        
        with open(output_path, 'rb') as f:
            docx_data = f.read()
        
        return docx_data
        
    finally:
        try:
            os.unlink(md_path)
            if os.path.exists(output_path):
                os.unlink(output_path)
        except:
            pass

# PYTHON-DOCX CONVERSION FUNCTIONS
def add_table_border(table):
    """Add borders to table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def parse_table(lines, start_idx):
    """Parse markdown table"""
    table_lines = []
    idx = start_idx
    
    while idx < len(lines) and '|' in lines[idx]:
        table_lines.append(lines[idx])
        idx += 1
    
    if len(table_lines) < 2:
        return None, start_idx
    
    headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
    rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
    
    return {'headers': headers, 'rows': rows}, idx

def convert_latex_to_unicode(latex_text):
    """Convert LaTeX to Unicode"""
    result = latex_text
    
    result = result.replace(r'\text{', '').replace(r'\mathrm{', '').replace(r'\mathbf{', '')
    result = result.replace(r'\hat{', '').replace(r'\bar{', '').replace(r'\tilde{', '')
    result = result.replace(r'\boxed{', '').replace(r'\left', '').replace(r'\right', '')
    result = result.replace(r'\begin{aligned}', '').replace(r'\end{aligned}', '')
    result = result.replace(r'\,', ' ').replace(r'\quad', '  ')
    
    replacements = {
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
        r'\epsilon': 'ε', r'\theta': 'θ', r'\lambda': 'λ', r'\mu': 'μ',
        r'\nu': 'ν', r'\pi': 'π', r'\rho': 'ρ', r'\sigma': 'σ',
        r'\tau': 'τ', r'\phi': 'φ', r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
        r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ',
        r'\Pi': 'Π', r'\Sigma': 'Σ', r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω',
        r'\times': '×', r'\div': '÷', r'\pm': '±', r'\cdot': '·',
        r'\leq': '≤', r'\geq': '≥', r'\neq': '≠', r'\approx': '≈', r'\equiv': '≡',
        r'\rightarrow': '→', r'\to': '→', r'\leftarrow': '←', r'\leftrightarrow': '↔',
        r'\Rightarrow': '⇒', r'\uparrow': '↑', r'\downarrow': '↓',
        r'\in': '∈', r'\subset': '⊂', r'\cup': '∪', r'\cap': '∩',
        r'\infty': '∞', r'\forall': '∀', r'\exists': '∃',
        r'\int': '∫', r'\sum': '∑', r'\prod': '∏', r'\partial': '∂', r'\nabla': '∇',
        r'\hbar': 'ℏ', r'\sqrt': '√', r'\angle': '∠',
    }
    
    for latex, unicode_char in replacements.items():
        result = result.replace(latex, unicode_char)
    
    result = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', result)
    
    def convert_superscript(match):
        sup_map = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹','+':'⁺','-':'⁻'}
        return ''.join(sup_map.get(c, c) for c in match.group(1))
    
    def convert_subscript(match):
        sub_map = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉','s':'ₛ','n':'ₙ','z':'ᵤ'}
        return ''.join(sub_map.get(c, c) for c in match.group(1))
    
    result = re.sub(r'\^\{([^}]+)\}', convert_superscript, result)
    result = re.sub(r'_\{([^}]+)\}', convert_subscript, result)
    
    result = result.replace('{', '').replace('}', '')
    result = re.sub(r'\\[a-zA-Z]+', '', result)
    result = result.replace('\\', '')
    
    return result.strip()

def extract_and_format_text(text, paragraph, font_size):
    """Format text with inline styles"""
    i = 0
    current_text = ""
    
    while i < len(text):
        processed = False
        
        if text[i:i+2] == '\\[':
            if current_text:
                paragraph.add_run(current_text).font.size = Pt(font_size)
                current_text = ""
            end = text.find('\\]', i + 2)
            if end != -1:
                math = convert_latex_to_unicode(text[i+2:end].strip())
                run = paragraph.add_run(math)
                run.font.name = 'Cambria Math'
                run.font.size = Pt(font_size + 1)
                run.font.color.rgb = RGBColor(0, 120, 0)
                run.bold = True
                i = end + 2
                processed = True
        
        if not processed and text[i:i+2] == '\\(':
            if current_text:
                paragraph.add_run(current_text).font.size = Pt(font_size)
                current_text = ""
            end = text.find('\\)', i + 2)
            if end != -1:
                math = convert_latex_to_unicode(text[i+2:end].strip())
                run = paragraph.add_run(' ' + math + ' ')
                run.font.name = 'Cambria Math'
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0, 120, 0)
                run.bold = True
                i = end + 2
                processed = True
        
        if not processed and text[i:i+2] == '**':
            if current_text:
                paragraph.add_run(current_text).font.size = Pt(font_size)
                current_text = ""
            end = text.find('**', i + 2)
            if end != -1:
                run = paragraph.add_run(text[i+2:end])
                run.bold = True
                run.font.size = Pt(font_size)
                i = end + 2
                processed = True
        
        if not processed and text[i] == '*' and (i == 0 or text[i-1] != '*') and (i+1 >= len(text) or text[i+1] != '*'):
            if current_text:
                paragraph.add_run(current_text).font.size = Pt(font_size)
                current_text = ""
            end = i + 1
            while end < len(text) and not (text[end] == '*' and (end+1 >= len(text) or text[end+1] != '*')):
                end += 1
            if end < len(text):
                run = paragraph.add_run(text[i+1:end])
                run.italic = True
                run.font.size = Pt(font_size)
                i = end + 1
                processed = True
        
        if not processed and text[i] == '`':
            if current_text:
                paragraph.add_run(current_text).font.size = Pt(font_size)
                current_text = ""
            end = text.find('`', i + 1)
            if end != -1:
                run = paragraph.add_run(text[i+1:end])
                run.font.name = 'Courier New'
                run.font.size = Pt(font_size - 1)
                run.font.color.rgb = RGBColor(220, 50, 50)
                i = end + 1
                processed = True
        
        if not processed:
            current_text += text[i]
            i += 1
    
    if current_text:
        paragraph.add_run(current_text).font.size = Pt(font_size)

def convert_with_python_docx(markdown_text, title, font_size, use_colors):
    """Convert with python-docx"""
    doc = Document()
    
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if '\\[' in markdown_text:
        markdown_text = re.sub(r'\\\[(.*?)\\\]', lambda m: '\\[' + m.group(1).replace('\n', ' ') + '\\]', markdown_text, flags=re.DOTALL)
    
    lines = markdown_text.split('\n')
    in_code = False
    code_lines = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        if '|' in line and i+1 < len(lines) and '---' in lines[i+1]:
            table_data, end_idx = parse_table(lines, i)
            if table_data:
                table = doc.add_table(rows=1+len(table_data['rows']), cols=len(table_data['headers']))
                table.style = 'Light Grid Accent 1'
                add_table_border(table)
                for idx, h in enumerate(table_data['headers']):
                    table.rows[0].cells[idx].text = h
                    for p in table.rows[0].cells[idx].paragraphs:
                        for r in p.runs:
                            r.font.bold = True
                            r.font.size = Pt(font_size)
                for r_idx, row in enumerate(table_data['rows']):
                    for c_idx, cell in enumerate(row):
                        if c_idx < len(table.rows[r_idx+1].cells):
                            table.rows[r_idx+1].cells[c_idx].text = cell
                i = end_idx
                continue
        
        if line.strip().startswith('```'):
            if in_code:
                para = doc.add_paragraph('\n'.join(code_lines))
                para.style = 'Intense Quote'
                for r in para.runs:
                    r.font.name = 'Courier New'
                    r.font.size = Pt(font_size - 1)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        
        if line.strip() == '---':
            para = doc.add_paragraph('─' * 80)
            for r in para.runs:
                r.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue
        
        if line.startswith('# ') and not line.startswith('## '):
            para = doc.add_heading('', 1)
            extract_and_format_text(line[2:], para, font_size + 2)
            if use_colors:
                for r in para.runs:
                    if not r.font.color.rgb:
                        r.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('## ') and not line.startswith('### '):
            para = doc.add_heading('', 2)
            extract_and_format_text(line[3:], para, font_size + 1)
            if use_colors:
                for r in para.runs:
                    if not r.font.color.rgb:
                        r.font.color.rgb = RGBColor(51, 102, 153)
        elif line.startswith('### '):
            para = doc.add_heading('', 3)
            extract_and_format_text(line[4:], para, font_size)
            if use_colors:
                for r in para.runs:
                    if not r.font.color.rgb:
                        r.font.color.rgb = RGBColor(102, 153, 204)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            extract_and_format_text(line.strip()[2:], para, font_size)
        elif re.match(r'^\d+\.\s', line.strip()):
            para = doc.add_paragraph(style='List Number')
            extract_and_format_text(re.sub(r'^\d+\.\s', '', line.strip()), para, font_size)
        elif line.strip().startswith('>'):
            para = doc.add_paragraph()
            para.style = 'Intense Quote'
            extract_and_format_text(line.strip()[1:].strip(), para, font_size)
        elif line.strip():
            para = doc.add_paragraph()
            extract_and_format_text(line, para, font_size)
        
        i += 1
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# Main content
st.subheader("📝 Paste Your Markdown Content")

if use_pandoc:
    st.info("**Math Syntax:** Use `$...$` for inline and `$$...$$` for display math (Pandoc format)")
else:
    st.info("**Math Syntax:** Use `\\(...\\)` for inline and `\\[...\\]` for display math (Python-docx format)")

markdown_input = st.text_area(
    "Markdown Content",
    value=st.session_state['markdown_content'],
    height=400,
    help="Paste your markdown here"
)

st.session_state['markdown_content'] = markdown_input

# Process button
if st.button("🔄 Process & Download", type="primary", use_container_width=True):
    with st.spinner("🔄 Processing..."):
        try:
            if use_pandoc:
                docx_data = convert_with_pandoc(
                    st.session_state['markdown_content'],
                    doc_title,
                    use_toc,
                    number_sections
                )
                st.success("✅ Converted with Pandoc - Equations are editable!")
            else:
                docx_data = convert_with_python_docx(
                    st.session_state['markdown_content'],
                    doc_title,
                    font_size,
                    use_colors
                )
                st.success("✅ Converted with Python-docx - Unicode symbols used")
            
            st.download_button(
                label="⬇️ Download Word Document",
                data=docx_data,
                file_name=f"{doc_title.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

# Footer
st.divider()
st.markdown(f"""
### 💡 Current Mode: **{'Pandoc (Editable Equations)' if use_pandoc else 'Python-docx (Unicode Symbols)'}**

**Pandoc Method:**
- ✅ Creates editable Word equations
- ✅ Professional quality
- ❌ Requires Pandoc installation
- ❌ Not available on Streamlit Cloud

**Python-docx Method:**
- ✅ Works everywhere (Streamlit Cloud, GitHub)
- ✅ No installation needed
- ❌ Unicode symbols only (not editable equations)
- ✅ Good for viewing/printing

**Recommendation:** Use Pandoc locally, Python-docx for cloud deployment
""")
